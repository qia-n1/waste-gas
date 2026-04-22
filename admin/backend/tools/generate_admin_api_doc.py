from __future__ import annotations

import json
import sys
from datetime import date
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


BASE_DIR = Path(__file__).resolve().parents[1]
ROOT_DIR = BASE_DIR.parents[1]
OUTPUT_PATH = ROOT_DIR / "admin" / "Admin端数据接口设计文档.docx"
BASE_URL = "http://localhost:8003"


def _load_openapi() -> dict[str, Any]:
    sys.path.insert(0, str(BASE_DIR))
    import main  # noqa: WPS433

    return main.app.openapi()


def _set_east_asia_font(run, font_name: str = "宋体", size: int | None = None, bold: bool | None = None) -> None:
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def _style_paragraph(paragraph, *, font_name: str = "宋体", size: int = 12, bold: bool = False) -> None:
    for run in paragraph.runs:
        _set_east_asia_font(run, font_name=font_name, size=size, bold=bold)


def _add_heading(doc: Document, text: str, level: int = 1) -> None:
    paragraph = doc.add_heading("", level=level)
    run = paragraph.add_run(text)
    _set_east_asia_font(run, font_name="黑体", size={1: 16, 2: 14, 3: 12}.get(level, 12), bold=True)


def _add_code_block(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.25)
    run = paragraph.add_run(text)
    _set_east_asia_font(run, font_name="Consolas", size=10)


def _json_text(data: Any) -> str:
    return json.dumps(data, ensure_ascii=False, indent=2)


def _schema_type(schema: dict[str, Any] | None) -> str:
    if not schema:
        return "-"
    if "$ref" in schema:
        return schema["$ref"].split("/")[-1]
    schema_type = schema.get("type")
    if schema_type == "array":
        items = schema.get("items", {})
        return f"array<{_schema_type(items)}>"
    if schema_type:
        return str(schema_type)
    any_of = schema.get("anyOf")
    if any_of:
        return " / ".join(_schema_type(item) for item in any_of)
    return "-"


def _add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        hdr_cells[idx].text = header
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                _style_paragraph(paragraph, size=10)


def _build_common_error_rows(path: str, method: str) -> list[list[str]]:
    rows = [["200", "调用成功", "返回业务数据"]]
    if path != "/api/auth/login":
        rows.append(["401", "未登录或 Token 无效", '{"detail":"未登录或凭证缺失"}'])
    if path in {"/api/alerts", "/api/users", "/api/users/{user_id}", "/api/users/{user_id}/reset-password", "/api/users/{user_id}/toggle-status", "/api/dashboard/anomaly-heatmap", "/api/dashboard/emitter-history/{emitter_id}", "/api/auth/login"}:
        rows.append(["422", "参数校验失败", '{"detail":[{"loc":["query","days"],"msg":"Input should be a valid integer","type":"int_parsing"}]}'])
    if path == "/api/users" and method == "post":
        rows.append(["409", "用户名已存在", '{"detail":"用户名已存在"}'])
    if path in {"/api/users/{user_id}", "/api/users/{user_id}/reset-password", "/api/users/{user_id}/toggle-status", "/api/dashboard/emitter-history/{emitter_id}"}:
        rows.append(["404", "资源不存在", '{"detail":"Unknown emitter: emitter-999"}'])
    return rows


ENDPOINTS: list[dict[str, Any]] = [
    {
        "name": "登录",
        "path": "/api/auth/login",
        "method": "post",
        "desc": "管理员登录，返回 Bearer Token 和当前用户信息。",
        "auth": "否",
        "request_example": {"username": "admin", "password": "admin123456"},
        "response_example": {
            "access_token": "eyJhbGciOiJIUzI1NiIs...",
            "token_type": "bearer",
            "user": {"username": "admin", "role": "SysAdmin", "name": "平台管理员"},
        },
        "response_fields": [
            ["access_token", "string", "JWT 访问令牌"],
            ["token_type", "string", "固定为 bearer"],
            ["user", "object", "当前登录用户信息"],
            ["user.username", "string", "用户名"],
            ["user.role", "string", "角色编码"],
            ["user.name", "string", "展示名称"],
        ],
    },
    {
        "name": "获取当前登录用户",
        "path": "/api/auth/me",
        "method": "get",
        "desc": "校验当前 Token 并返回简要用户信息。",
        "auth": "是",
        "request_example": None,
        "response_example": {"user": {"username": "admin", "role": "SysAdmin"}},
        "response_fields": [
            ["user", "object", "当前用户信息"],
            ["user.username", "string", "用户名"],
            ["user.role", "string", "角色编码"],
        ],
    },
    {
        "name": "获取大屏总览",
        "path": "/api/dashboard/overview",
        "method": "get",
        "desc": "返回管理端首页大屏所需的总览数据，包括指标、趋势、告警、工厂节点和决策建议。",
        "auth": "是",
        "request_example": None,
        "response_example": {
            "timestamp": "2026-04-21T09:30:00+08:00",
            "metrics": {
                "currentVocs": 56.2,
                "peakForecast": 92.5,
                "alertLevel": "warning",
                "onlineDevices": 148,
                "totalDevices": 150,
                "todayAlerts": 3,
                "systemPhase": "联机运行",
                "uptime": "12h 31m",
                "confidence": 88,
                "dataCompleteness": 100,
                "latencyMs": 280,
                "predictionType": "DLinear-PCA-Ensemble",
            },
            "trend": {
                "actualSeries": [{"timestamp": "2026-04-21T08:00:00+08:00", "value": 42.1}],
                "forecastSeries": [{"timestamp": "2026-04-21T09:45:00+08:00", "value": 75.6}],
                "warningThreshold": 80.0,
                "criticalThreshold": 100.0,
                "confidence": 0.88,
            },
            "statusBanner": {"severity": "warning", "text": "VOCs 当前 56.2 mg/m³，接近橙色预警阈值，请重点关注排口状态。"},
            "keyParameters": [{"field": "rto_out_conc", "label": "RTO出口浓度", "value": 56.2, "unit": "mg/m³"}],
            "decision": {"summary": "当前 RTO 出口浓度为 56.2 mg/m³。", "suggestions": ["建议提前安排巡检。"]},
            "continuousAlerts": [{"id": "ALT-1", "level": "warning", "message": "1号排口浓度持续偏高", "location": "1号排口", "elapsed_seconds": 35}],
            "factoryNodes": [{"id": "stack", "label": "1号排口", "status": "warning", "x": 72, "y": 45}],
        },
        "response_fields": [
            ["timestamp", "string", "总览生成时间"],
            ["metrics", "object", "顶部统计指标"],
            ["trend", "object", "VOCs 趋势与预测"],
            ["statusBanner", "object", "状态提示条"],
            ["keyParameters", "array<object>", "关键参数列表"],
            ["decision", "object", "AI 决策建议"],
            ["continuousAlerts", "array<object>", "异常持续关注列表"],
            ["factoryNodes", "array<object>", "工厂场景点位"],
            ["attribution", "object", "可选，模型归因信息"],
            ["topContributorSeries", "array<object>", "可选，关键贡献因子序列"],
        ],
    },
    {
        "name": "获取设备状态分布",
        "path": "/api/dashboard/equipment-status",
        "method": "get",
        "desc": "返回设备状态环图数据。",
        "auth": "是",
        "request_example": None,
        "response_example": {
            "total": 150,
            "online": 148,
            "items": [
                {"name": "正常", "value": 132, "color": "#46d1ff"},
                {"name": "预警", "value": 11, "color": "#ffb347"},
                {"name": "故障", "value": 5, "color": "#ff5b61"},
                {"name": "离线", "value": 2, "color": "#5f6d95"},
            ],
        },
        "response_fields": [
            ["total", "integer", "设备总数"],
            ["online", "integer", "在线设备数"],
            ["items", "array<object>", "分类统计项"],
            ["items[].name", "string", "分类名称"],
            ["items[].value", "integer", "分类数量"],
            ["items[].color", "string", "前端展示颜色"],
        ],
    },
    {
        "name": "获取异常热力图",
        "path": "/api/dashboard/anomaly-heatmap",
        "method": "get",
        "desc": "返回按日期和小时聚合的异常热力图数据。",
        "auth": "是",
        "request_example": None,
        "response_example": {
            "dates": ["04/15", "04/16", "04/17"],
            "hours": [0, 1, 2, 3],
            "values": [[0, 8, 1], [0, 9, 2], [1, 14, 3]],
        },
        "response_fields": [
            ["dates", "array<string>", "横轴日期列表"],
            ["hours", "array<integer>", "纵轴小时列表"],
            ["values", "array<array>", "热力值，格式为 [日期索引, 小时, 权重]"],
        ],
    },
    {
        "name": "获取排口历史",
        "path": "/api/dashboard/emitter-history/{emitter_id}",
        "method": "get",
        "desc": "返回单个排口最近一段时间的历史数据，用于工厂场景弹窗。",
        "auth": "是",
        "request_example": None,
        "response_example": {
            "emitterId": "stack-1",
            "label": "1号排口",
            "limit": 48,
            "series": [
                {"timestamp": "2026-04-21T08:00:00+08:00", "value": 48.6},
                {"timestamp": "2026-04-21T08:15:00+08:00", "value": 52.1},
            ],
        },
        "response_fields": [
            ["emitterId", "string", "排口编码"],
            ["label", "string", "排口名称"],
            ["limit", "integer", "返回点数上限"],
            ["series", "array<object>", "历史序列"],
            ["series[].timestamp", "string", "采集时间"],
            ["series[].value", "number", "数值"],
        ],
    },
    {
        "name": "触发集成模型预测",
        "path": "/api/dashboard/predict",
        "method": "post",
        "desc": "触发一次管理端预测聚合；若集成模型不可用则返回错误状态。",
        "auth": "是",
        "request_example": None,
        "response_example": {
            "status": "success",
            "predictions": [62.8, 64.2, 66.9],
            "is_exceed_warning": True,
            "incremental_attribution": {
                "baseline": 51.2,
                "target": 66.9,
                "total_increment": 15.7,
            },
        },
        "response_fields": [
            ["status", "string", "success 或 error"],
            ["predictions", "array<number>", "未来序列预测值"],
            ["is_exceed_warning", "boolean", "是否超过预警阈值"],
            ["incremental_attribution", "object", "可选，增量归因结果"],
        ],
    },
    {
        "name": "检查集成模型健康状态",
        "path": "/api/dashboard/ensemble-health",
        "method": "get",
        "desc": "探测集成模型服务是否可达。",
        "auth": "是",
        "request_example": None,
        "response_example": {"connected": True, "status": "ok", "service": "ensemble-api"},
        "response_fields": [
            ["connected", "boolean", "模型服务是否连通"],
            ["status", "string", "模型服务状态"],
            ["service", "string", "服务名称"],
        ],
    },
    {
        "name": "获取告警列表",
        "path": "/api/alerts",
        "method": "get",
        "desc": "返回实时告警中心列表，包含共享告警和 admin 端设备离线 watchdog 告警。",
        "auth": "是",
        "request_example": None,
        "response_example": {
            "items": [
                {
                    "alert_id": "WATCHDOG-1713661200000",
                    "timestamp": "2026-04-21T09:40:00+08:00",
                    "level": "critical",
                    "message": "设备数据采集中断 95 秒，超过 90s 阈值，请立即检查传感器与上行网关。",
                    "value": 0.0,
                    "threshold": 90.0,
                    "acknowledged": False,
                    "location": "传感器网关",
                    "status": "处理中",
                }
            ],
            "total": 1,
            "byLevel": {"critical": 1},
        },
        "response_fields": [
            ["items", "array<object>", "告警列表"],
            ["items[].alert_id", "string", "告警唯一标识"],
            ["items[].timestamp", "string", "告警时间"],
            ["items[].level", "string", "等级：info/warning/critical"],
            ["items[].message", "string", "告警文案"],
            ["items[].value", "number", "当前告警值"],
            ["items[].threshold", "number", "阈值"],
            ["items[].acknowledged", "boolean", "是否已确认"],
            ["items[].location", "string", "位置"],
            ["items[].status", "string", "处理状态"],
            ["total", "integer", "总条数"],
            ["byLevel", "object", "按等级聚合计数"],
        ],
    },
    {
        "name": "确认告警",
        "path": "/api/alerts/{alert_id}/acknowledge",
        "method": "post",
        "desc": "确认指定告警。watchdog 告警在 admin 内存中本地确认，其它告警转发给共享服务。",
        "auth": "是",
        "request_example": None,
        "response_example": {"success": True, "message": "设备状态告警已确认", "alert_id": "WATCHDOG-1713661200000"},
        "response_fields": [
            ["success", "boolean", "是否确认成功"],
            ["message", "string", "确认结果说明"],
            ["alert_id", "string", "告警标识"],
        ],
    },
    {
        "name": "获取告警诊断",
        "path": "/api/alerts/{alert_id}/diagnosis",
        "method": "get",
        "desc": "获取单条告警的诊断摘要、建议与贡献因子。",
        "auth": "是",
        "request_example": None,
        "response_example": {
            "alertId": "ALT-1001",
            "summary": "当前 RTO 出口浓度上升明显，建议检查焚烧温度与喷涂浓度。",
            "recommendations": ["建议提前安排巡检", "保持燃烧温度稳定"],
            "contributors": [{"label": "RTO出口浓度", "group": "", "weight": 0.34, "contribution": 2.1}],
            "groupContributions": [{"group": "RTO焚烧系统", "contribution": 5.2}],
            "baseline": 51.2,
            "target": 66.9,
            "totalIncrement": 15.7,
        },
        "response_fields": [
            ["alertId", "string", "告警标识"],
            ["summary", "string", "诊断摘要"],
            ["recommendations", "array<string>", "处理建议列表"],
            ["contributors", "array<object>", "贡献因子列表"],
            ["groupContributions", "array<object>", "分组贡献结果"],
            ["baseline", "number", "基线值"],
            ["target", "number", "目标值"],
            ["totalIncrement", "number", "增量值"],
        ],
    },
    {
        "name": "获取用户列表",
        "path": "/api/users",
        "method": "get",
        "desc": "根据关键字、角色、状态筛选用户列表。",
        "auth": "是",
        "request_example": None,
        "response_example": {
            "items": [
                {
                    "id": 1,
                    "username": "admin_user",
                    "display_name": "超级管理员",
                    "role_code": "SysAdmin",
                    "role_name": "超级管理员",
                    "status": "enabled",
                    "created_at": "2024-01-01 10:00",
                    "last_login_at": "2026-04-13 09:30",
                }
            ],
            "total": 1,
            "roles": [
                {"code": "SysAdmin", "name": "超级管理员"},
                {"code": "EnvAdmin", "name": "环保监测员"},
            ],
        },
        "response_fields": [
            ["items", "array<object>", "用户列表"],
            ["items[].id", "integer", "用户 ID"],
            ["items[].username", "string", "用户名"],
            ["items[].display_name", "string", "姓名/展示名"],
            ["items[].role_code", "string", "角色编码"],
            ["items[].role_name", "string", "角色中文名"],
            ["items[].status", "string", "状态：enabled/disabled"],
            ["created_at / last_login_at", "string", "时间字段，格式 yyyy-MM-dd HH:mm"],
            ["total", "integer", "记录数"],
            ["roles", "array<object>", "可选角色列表"],
        ],
    },
    {
        "name": "新增用户",
        "path": "/api/users",
        "method": "post",
        "desc": "创建用户并返回新建的用户对象。",
        "auth": "是",
        "request_example": {
            "username": "env_admin_02",
            "display_name": "环保管理员二号",
            "role_code": "EnvAdmin",
            "status": "enabled",
            "password": "Admin@123",
        },
        "response_example": {
            "item": {
                "id": 5,
                "username": "env_admin_02",
                "display_name": "环保管理员二号",
                "role_code": "EnvAdmin",
                "role_name": "环保监测员",
                "status": "enabled",
                "created_at": "2026-04-21 10:10",
                "last_login_at": "--",
            }
        },
        "response_fields": [
            ["item", "object", "新建后的用户对象"],
            ["item.id", "integer", "用户 ID"],
            ["item.username", "string", "用户名"],
            ["item.role_code / role_name", "string", "角色信息"],
        ],
    },
    {
        "name": "获取用户详情",
        "path": "/api/users/{user_id}",
        "method": "get",
        "desc": "获取单个用户详情及菜单权限。",
        "auth": "是",
        "request_example": None,
        "response_example": {
            "item": {
                "id": 1,
                "username": "admin_user",
                "display_name": "超级管理员",
                "role_code": "SysAdmin",
                "role_name": "超级管理员",
                "status": "enabled",
                "created_at": "2024-01-01 10:00",
                "last_login_at": "2026-04-13 09:30",
            },
            "permissions": {"menus": ["大屏", "监测", "预测", "告警", "模型", "用户", "日志", "部署"]},
        },
        "response_fields": [
            ["item", "object", "用户基本信息"],
            ["permissions", "object", "权限集合"],
            ["permissions.menus", "array<string>", "可访问菜单"],
        ],
    },
    {
        "name": "更新用户",
        "path": "/api/users/{user_id}",
        "method": "put",
        "desc": "更新用户展示名、角色和状态。",
        "auth": "是",
        "request_example": {"display_name": "环保管理员", "role_code": "EnvAdmin", "status": "enabled"},
        "response_example": {
            "item": {
                "id": 2,
                "username": "env_monitor",
                "display_name": "环保管理员",
                "role_code": "EnvAdmin",
                "role_name": "环保监测员",
                "status": "enabled",
                "created_at": "2024-02-15 08:00",
                "last_login_at": "2026-04-13 09:00",
            }
        },
        "response_fields": [
            ["item", "object", "更新后的用户对象"],
            ["item.status", "string", "更新后的启用状态"],
        ],
    },
    {
        "name": "重置用户密码",
        "path": "/api/users/{user_id}/reset-password",
        "method": "post",
        "desc": "重置指定用户密码，初始密码固定为 Reset@123。",
        "auth": "是",
        "request_example": None,
        "response_example": {"success": True, "message": "已为 超级管理员 重置密码，初始密码为 Reset@123"},
        "response_fields": [
            ["success", "boolean", "是否重置成功"],
            ["message", "string", "结果说明"],
        ],
    },
    {
        "name": "切换用户启用状态",
        "path": "/api/users/{user_id}/toggle-status",
        "method": "post",
        "desc": "在 enabled 与 disabled 间切换用户状态。",
        "auth": "是",
        "request_example": None,
        "response_example": {
            "item": {
                "id": 3,
                "username": "data_analyst_01",
                "display_name": "数据分析师",
                "role_code": "Analyst",
                "role_name": "数据分析师",
                "status": "enabled",
                "created_at": "2025-06-20 11:30",
                "last_login_at": "2026-04-10 16:45",
            }
        },
        "response_fields": [
            ["item", "object", "切换后的用户对象"],
            ["item.status", "string", "切换后的状态"],
        ],
    },
    {
        "name": "设备告警 SSE 流",
        "path": "/api/events/device-alerts",
        "method": "get",
        "desc": "Server-Sent Events 流接口。设备 90 秒未收到新数据或恢复时，向前端实时推送 device_alert 事件。",
        "auth": "否",
        "request_example": None,
        "response_example": "event: connected\\ndata: {\"online\":true,\"lastSeen\":\"2026-04-21T09:30:00+08:00\",\"elapsedSeconds\":3.1,\"timeoutThreshold\":90}\\n\\nevent: device_alert\\ndata: {\"alert_id\":\"WATCHDOG-1713661200000\",\"level\":\"critical\",\"message\":\"设备数据采集中断 95 秒，超过 90s 阈值，请立即检查传感器与上行网关。\"}\\n",
        "response_fields": [
            ["event", "string", "事件名，connected 或 device_alert"],
            ["data", "string/json", "事件负载 JSON 文本"],
        ],
    },
    {
        "name": "获取设备监控状态",
        "path": "/api/events/device-status",
        "method": "get",
        "desc": "返回 watchdog 当前状态，用于诊断设备离线监控功能。",
        "auth": "否",
        "request_example": None,
        "response_example": {
            "online": True,
            "lastSeen": "2026-04-21T09:30:00+08:00",
            "elapsedSeconds": 12.4,
            "timeoutThreshold": 90,
            "checkInterval": 10,
            "bufferedAlerts": 2,
        },
        "response_fields": [
            ["online", "boolean", "当前是否在线"],
            ["lastSeen", "string|null", "最近一次心跳时间"],
            ["elapsedSeconds", "number|null", "距最近一次心跳的秒数"],
            ["timeoutThreshold", "integer", "超时时间阈值，单位秒"],
            ["checkInterval", "integer", "轮询间隔，单位秒"],
            ["bufferedAlerts", "integer", "缓存中的 watchdog 告警数量"],
        ],
    },
    {
        "name": "健康检查",
        "path": "/api/health",
        "method": "get",
        "desc": "管理端后端服务健康检查。",
        "auth": "否",
        "request_example": None,
        "response_example": {"status": "ok", "service": "Waste Gas Admin API"},
        "response_fields": [
            ["status", "string", "固定为 ok"],
            ["service", "string", "服务名称"],
        ],
    },
]


def _build_request_params(path_item: dict[str, Any], endpoint: dict[str, Any]) -> list[list[str]]:
    rows: list[list[str]] = []
    for param in path_item.get("parameters", []):
        schema = param.get("schema", {})
        rows.append(
            [
                param.get("name", "-"),
                param.get("in", "-"),
                _schema_type(schema),
                "是" if param.get("required") else "否",
                "-" if "default" not in schema else str(schema.get("default")),
                endpoint.get("param_desc", {}).get(param.get("name", ""), "-"),
            ]
        )
    return rows


def _build_request_body_rows(path_item: dict[str, Any], schemas: dict[str, Any], endpoint: dict[str, Any]) -> list[list[str]]:
    request_body = path_item.get("requestBody", {})
    content = request_body.get("content", {}).get("application/json", {})
    schema = content.get("schema")
    if not schema:
        return []
    if "$ref" not in schema:
        return [["body", "object", "是", "-", "请求体"]]
    schema_name = schema["$ref"].split("/")[-1]
    schema_obj = schemas.get(schema_name, {})
    required = set(schema_obj.get("required", []))
    rows: list[list[str]] = []
    for field_name, field_schema in schema_obj.get("properties", {}).items():
        desc = endpoint.get("body_desc", {}).get(field_name, "-")
        constraints: list[str] = []
        if "minLength" in field_schema:
            constraints.append(f"minLength={field_schema['minLength']}")
        if "maxLength" in field_schema:
            constraints.append(f"maxLength={field_schema['maxLength']}")
        rows.append(
            [
                field_name,
                _schema_type(field_schema),
                "是" if field_name in required else "否",
                ", ".join(constraints) if constraints else "-",
                desc,
            ]
        )
    return rows


def _curl_example(endpoint: dict[str, Any]) -> str:
    method = endpoint["method"].upper()
    path = endpoint["path"]
    auth_header = '  -H "Authorization: Bearer <token>" \\\n' if endpoint["auth"] == "是" else ""
    json_body = endpoint.get("request_example")
    if method == "GET":
        return f'curl -X {method} "{BASE_URL}{path}" \\\n{auth_header}'.rstrip(" \\\n")
    body_text = ""
    if json_body is not None:
        body_text = f' \\\n  -H "Content-Type: application/json" \\\n  -d \'{json.dumps(json_body, ensure_ascii=False)}\''
    return f'curl -X {method} "{BASE_URL}{path}" \\\n{auth_header}{body_text}'.rstrip()


def main() -> None:
    spec = _load_openapi()
    schemas = spec.get("components", {}).get("schemas", {})

    doc = Document()
    section = doc.sections[0]
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Admin端数据接口设计文档")
    _set_east_asia_font(run, font_name="黑体", size=18, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(f"版本：v1.0    生成日期：{date.today().isoformat()}")
    _set_east_asia_font(run, size=11)

    _add_heading(doc, "1. 文档说明", level=1)
    p = doc.add_paragraph("本文档面向废气综合管理平台 Admin 端前后端联调与验收，采用 OpenAPI 风格组织，覆盖接口参数、调用示例、数据格式说明与错误码定义。")
    _style_paragraph(p)
    p = doc.add_paragraph(f"本地开发默认服务地址：{BASE_URL}；前端 Vite 开发代理通过 /api 转发至该服务。")
    _style_paragraph(p)

    _add_heading(doc, "2. 通用约定", level=1)
    for line in [
        "协议：HTTP/HTTPS",
        "数据格式：请求体与响应体默认使用 application/json；SSE 接口使用 text/event-stream",
        "字符编码：UTF-8",
        "认证方式：除登录、设备状态流、设备状态快照、健康检查外，其它接口均通过 Bearer Token 认证",
        "认证头示例：Authorization: Bearer <token>",
    ]:
        p = doc.add_paragraph(line, style=None)
        _style_paragraph(p)

    _add_heading(doc, "3. 通用错误码", level=1)
    _add_table(
        doc,
        ["HTTP状态码", "含义", "说明"],
        [
            ["200", "成功", "接口调用成功"],
            ["401", "未认证", "Token 缺失、无效或已过期"],
            ["404", "资源不存在", "用户、排口等资源不存在"],
            ["409", "冲突", "创建资源时主键或用户名重复"],
            ["422", "参数校验失败", "请求参数类型或格式不合法"],
            ["500", "服务内部错误", "后端未捕获异常或上游依赖异常"],
        ],
    )

    _add_heading(doc, "4. 接口清单", level=1)
    summary_rows = []
    for endpoint in ENDPOINTS:
        summary_rows.append([endpoint["name"], endpoint["method"].upper(), endpoint["path"], endpoint["auth"], endpoint["desc"]])
    _add_table(doc, ["接口名称", "方法", "路径", "鉴权", "说明"], summary_rows)

    _add_heading(doc, "5. 接口详细设计", level=1)
    for index, endpoint in enumerate(ENDPOINTS, start=1):
        path_item = spec["paths"][endpoint["path"]][endpoint["method"]]
        _add_heading(doc, f"5.{index} {endpoint['name']}", level=2)

        basic_rows = [
            ["接口路径", endpoint["path"]],
            ["请求方法", endpoint["method"].upper()],
            ["是否鉴权", endpoint["auth"]],
            ["接口说明", endpoint["desc"]],
        ]
        _add_table(doc, ["项", "内容"], basic_rows)

        param_rows = _build_request_params(path_item, endpoint)
        if param_rows:
            p = doc.add_paragraph("请求参数")
            _style_paragraph(p, bold=True)
            _add_table(doc, ["参数名", "位置", "类型", "必填", "默认值", "说明"], param_rows)

        body_rows = _build_request_body_rows(path_item, schemas, endpoint)
        if body_rows:
            p = doc.add_paragraph("请求体")
            _style_paragraph(p, bold=True)
            _add_table(doc, ["字段", "类型", "必填", "约束", "说明"], body_rows)

        p = doc.add_paragraph("调用示例")
        _style_paragraph(p, bold=True)
        _add_code_block(doc, _curl_example(endpoint))

        if endpoint.get("request_example") is not None:
            p = doc.add_paragraph("请求示例")
            _style_paragraph(p, bold=True)
            _add_code_block(doc, _json_text(endpoint["request_example"]))

        p = doc.add_paragraph("成功响应示例")
        _style_paragraph(p, bold=True)
        response_example = endpoint["response_example"]
        _add_code_block(doc, response_example if isinstance(response_example, str) else _json_text(response_example))

        p = doc.add_paragraph("数据格式说明")
        _style_paragraph(p, bold=True)
        _add_table(doc, ["字段", "类型", "说明"], endpoint["response_fields"])

        p = doc.add_paragraph("错误码定义")
        _style_paragraph(p, bold=True)
        _add_table(doc, ["状态码", "含义", "示例"], _build_common_error_rows(endpoint["path"], endpoint["method"]))

    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT_PATH))
    print(OUTPUT_PATH)


if __name__ == "__main__":
    main()
